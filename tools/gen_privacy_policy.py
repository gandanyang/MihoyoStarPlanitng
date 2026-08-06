# -*- coding: utf-8 -*-
"""《归星物语》隐私政策生成脚本

从模板 privacy-policy-template.docx 生成项目定制版：
docs/legal/归星物语隐私政策-v0.1.docx

用法：
    python tools/gen_privacy_policy.py

注意：主体名称 / 注册地址 / 联系方式为【待补充】占位，
信息确定后修改下方 CONFIG 再重新运行即可。
"""

import re
from pathlib import Path

import docx

ROOT = Path(__file__).resolve().parent.parent
TEMPLATE = ROOT / "docs" / "legal" / "privacy-policy-template-source.docx"
OUTPUT = ROOT / "docs" / "legal" / "归星物语隐私政策-v0.1.docx"

# ============ 可编辑配置（确定后修改这里） ============
CONFIG = {
    "game_name": "归星物语",
    "operator_name": "归星物语开发组",          # 【待补充】正式运营主体名称
    "operator_address": "【请补充运营主体注册地址】",  # 【待补充】
    "contact_email": "【待补充】",               # 【待补充】客服邮箱
    "contact_phone": "【待补充】",               # 【待补充】联系电话
    "update_date": "2026年8月5日",
}
# ====================================================


def set_para_text(para, new_text: str):
    """保留段落样式，整体替换文本（合并多个 run）。"""
    if para.runs:
        para.runs[0].text = new_text
        for r in para.runs[1:]:
            r.text = ""
    else:
        para.add_run(new_text)


def replace_para(para, mapping: dict) -> bool:
    full = para.text
    changed = full
    for k, v in mapping.items():
        changed = changed.replace(k, v)
    if changed != full:
        set_para_text(para, changed)
        return True
    return False


def rewrite_body(para, keyword: str, new_text: str) -> bool:
    """标题段保留，正文段（含 keyword 且比标题长）替换为新正文。"""
    t = para.text.strip()
    if t and keyword in para.text and len(t) > len(keyword):
        set_para_text(para, new_text)
        return True
    return False


def rename_heading(para, old: str, new: str) -> bool:
    if para.text.strip() == old:
        set_para_text(para, new)
        return True
    return False


def delete_para(para):
    el = para._element
    el.getparent().remove(el)


def rewrite_ellipsis(para, last_heading: str) -> bool:
    """模板中的省略段（……）按上一个标题改写。"""
    t = para.text.strip()
    if not t:
        return False
    if t.replace("…", "").strip() != "":
        return False
    mapping = {
        "充值消费": "当前版本未提供充值消费功能，不收集任何支付信息。",
        "不使用 Cookie 与跟踪技术": "（当前版本不使用上述 Cookie。）",
        "本地存档说明": "游戏存档仅保存在您的设备本地，不会同步到任何服务器。",
    }
    if last_heading in mapping:
        set_para_text(para, mapping[last_heading])
        return True
    return False


def delete_row(row):
    tr = row._tr
    tr.getparent().remove(tr)


def main():
    doc = docx.Document(TEMPLATE)
    cfg = CONFIG
    last_heading = ""
    in_toc = False

    # ---------- 全局精确替换 ----------
    global_map = {
        "【游戏产品名称】": cfg["game_name"],
        "【运营主体名称】": cfg["operator_name"],
        "【运营者名称】": cfg["operator_name"],
        "【注册地址详细信息】": cfg["operator_address"],
        "【设备权限名称】": "网络",
        "【电话号码】": cfg["contact_phone"],
        "【隐私政策连接】": "—",
    }

    # ---------- 段落级处理 ----------
    for para in doc.paragraphs:
        replace_para(para, global_map)

        # 标题
        if re.fullmatch(r"【\s*】隐私政策", para.text):
            set_para_text(para, f"{cfg['game_name']}隐私政策")
        elif re.fullmatch(r"【\s*】", para.text):
            set_para_text(para, f"{cfg['game_name']}隐私政策")

        # 小节标题重命名
        if rename_heading(para, "我们如何使用Cookies", "我们如何使用本地存储" if in_toc else "不使用 Cookie 与跟踪技术"):
            last_heading = "不使用 Cookie 与跟踪技术"
            continue
        if rename_heading(para, "我们如何使用Cookies和同类技术", "我们如何使用本地存储"):
            last_heading = "我们如何使用本地存储"
            continue
        if rename_heading(para, "我们如何使用同类技术", "本地存档说明"):
            last_heading = "本地存档说明"
            continue

        # 目录区标记：目录最后一项是"如何联系我们"
        if para.text.strip() == "本隐私政策将帮助您了解以下内容：":
            in_toc = True
        elif para.text.strip() == "如何联系我们" and in_toc:
            in_toc = False

        # 日期
        if "2021年" in para.text and "月" in para.text and "日" in para.text:
            set_para_text(
                para,
                re.sub(r"2021年【\s*】月【\s*】日", cfg["update_date"], para.text),
            )

        # 空占位（访问/更正路径等）→ 联系方式
        if re.search(r"【\s*】", para.text):
            set_para_text(
                para,
                re.sub(r"【\s*】", "本隐私政策载明的联系方式", para.text),
            )

        # ---------- 正文段按项目实际改写（标题保留） ----------
        rewrite_body(
            para,
            "当您注册",
            "当前版本未提供账号注册功能，您无需注册即可使用游戏，"
            "我们不会收集您的手机号码、登录密码等注册信息。",
        )
        rewrite_body(
            para,
            "设备ID",
            "当您使用游戏服务时，您的游戏进度与存档仅保存在您的设备本地"
            "（浏览器本地存储），我们不会收集、上传您的游戏数据或设备信息。",
        )
        rewrite_body(
            para,
            "身份证件",
            "当前版本未提供实名认证功能，不收集您的姓名、身份证件等信息。"
            "如后续版本依据相关法规接入实名认证，我们将依照法律规定收集必要信息，"
            "并及时更新本政策。",
        )
        rewrite_body(
            para,
            "设备识别符",
            "当前版本为本地单机游戏，无联网对战或在线服务，"
            "我们不收集您的设备识别符、IP地址、访问日期和时间等信息。",
        )
        rewrite_body(
            para,
            "委托处理",
            "当前版本未委托任何第三方处理您的个人信息。",
        )
        rewrite_body(
            para,
            "业务合作伙伴",
            "当前版本无账号系统、无广告及统计合作，"
            "我们不向任何业务合作伙伴共享您的个人信息。",
        )
        rewrite_body(
            para,
            "通过使用Cookies",
            "本游戏为本地单机运行，不使用 Cookie 或任何跟踪技术。"
            "游戏存档通过浏览器本地存储（localStorage）保存在您的设备上，"
            "由您自行管理，我们不会读取或上传该数据。",
        )
        rewrite_body(
            para,
            "第三方合作伙伴通过Cookies",
            "当前版本未接入第三方跟踪技术，"
            "不存在第三方通过 Cookie 或同类技术收集您信息的情形。",
        )
        rewrite_body(
            para,
            "aboutcookies",
            "如果您的浏览器允许，您可以在浏览器设置中管理或清除本地存储数据；"
            "清除后游戏存档可能丢失，请您知悉。",
        )
        rewrite_body(
            para,
            "第三方软件开发工具包",
            "当前版本未接入任何第三方广告、统计或数据分析 SDK。"
            "我们使用 Capacitor 将游戏打包为移动应用，其在本机系统层运行，"
            "不收集您的个人信息。",
        )
        rewrite_body(
            para,
            "如您希望访问或更正",
            "当前版本无账号系统，游戏中不存在可访问或更正的注册类个人信息；"
            "如您对本地存档有疑问，可通过本政策载明的方式与我们联系。",
        )
        rewrite_body(
            para,
            "如果您希望注销",
            "当前版本未提供账户功能，无需注销；"
            "如您希望删除本地存档数据，可卸载应用或在应用管理中清除数据。",
        )
        rewrite_body(
            para,
            "书面疑问",
            f"如果您对本政策或个人信息保护有任何问题，可以通过以下方式与我们联系：\n"
            f"名称：{cfg['operator_name']}\n"
            f"地址：{cfg['operator_address']}\n"
            f"联系邮箱：{cfg['contact_email']}\n"
            f"联系电话：{cfg['contact_phone']}",
        )

        # 省略段按上一标题改写
        rewrite_ellipsis(para, last_heading)

        # 删除模板残留的旧联系方式行与 Cookies 列表项
        t = para.text.strip()
        if (
            t.startswith("安全类Cookies：")
            or t.startswith("推荐类Cookies：")
            or t.startswith("名称：")
            or t.startswith("地址：")
            or t.startswith("或您也可以通过")
            or t.startswith("联系电话：")
            or t.startswith("您注销上述账户的行为")
        ):
            delete_para(para)
            continue

        # 记录小节标题（用于省略段定位）
        if t and len(t) <= 20 and not t.startswith("（") and "：" not in t[:4]:
            last_heading = t

    # ---------- 表格处理 ----------
    if doc.tables:
        perm_table = doc.tables[0]  # 设备权限表
        if len(perm_table.rows) > 1:
            # 保留表头 + 网络行；删除麦克风行与省略行
            for row in list(perm_table.rows[2:]):
                delete_row(row)
            # 网络行目的描述
            for cell in perm_table.rows[1].cells:
                for p in cell.paragraphs:
                    replace_para(p, {"连接网络": "保障游戏基础功能运行（当前版本为本地单机，权限预留）"})

        sdk_table = doc.tables[1]  # 第三方 SDK 表
        if len(sdk_table.rows) > 1:
            for row in list(sdk_table.rows[2:]):
                delete_row(row)
            cells = sdk_table.rows[1].cells
            sdk_texts = ["无", "当前版本未接入第三方 SDK", "不收集", "—"]
            for i, cell in enumerate(cells):
                if i < len(sdk_texts):
                    set_para_text(cell.paragraphs[0], sdk_texts[i])

    doc.save(OUTPUT)
    print(f"已生成: {OUTPUT}")


if __name__ == "__main__":
    main()
